from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

output_path = r"C:\Users\arika\OneDrive\Litigation\Litigation Downloads\Writing test 2\Style Comparison - PSJ Firearms Conversion Count IV - 2026-03-29.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

FONT_NAME = "Aptos"

def set_run(run, size, bold=False, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = FONT_NAME

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, 16, bold=True)
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, 11)
    p.paragraph_format.space_after = Pt(6)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, 13, bold=True)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, 11, bold=True)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        set_run(run, 11)
    return p

def add_edge(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, 11, italic=True)
    p.paragraph_format.space_after = Pt(10)

def add_summary_table(doc, rows):
    doc.add_paragraph()
    add_section_heading(doc, "Summary")
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Dimension", "Edge"]):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = FONT_NAME
    for i, (dim, edge) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for cell, text in zip(cells, [dim, edge]):
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(11)
            run.font.name = FONT_NAME

# ---- TITLE ----
add_title(doc, "Legal Style Comparison")
add_subtitle(doc, "PSJ Motion \u2014 Firearms Conversion, Count IV  |  2026-03-29")

# ---- DOCUMENT KEY ----
filenames = {
    "A": "PSJ MotionFirearms Conversion Count IV-rewritten -A.docx",
    "B": "Doc1_PSJ_Motion_Firearms_Count_IV_REWRITTEN-B.docx",
    "C": "PSJ Motion - Firearms Conversion (Rewritten- C).docx",
}
for label, fname in filenames.items():
    p = doc.add_paragraph()
    r = p.add_run(f"Document {label}: {fname}")
    r.font.size = Pt(11)
    r.font.name = FONT_NAME
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()  # spacer

# ---- SECTION 1 ----
add_section_heading(doc, "Dimension-by-Dimension Analysis")

add_sub_heading(doc, "1. Conciseness vs. Wordiness")
add_body(doc, "All three documents are reasonably concise for summary judgment motions, but they diverge at the margins.")
add_body(doc, 'Document A opens its Introduction with: "As set forth below, the undisputed facts established by documentary evidence demonstrate that..." \u2014 "As set forth below" adds nothing; it is empty procedural preamble. The WHEREFORE header is labeled "Relief Requested / WHEREFORE Clause" \u2014 a redundant double-label that reflects formatting uncertainty rather than intention.')
add_body(doc, 'Document B trims the same passage: "The undisputed facts, established by documentary evidence, demonstrate that..." \u2014 three words shorter and no filler phrase. Its conversion element section also tightens one passage: "Ms. Lavigne has since refused..." (B) vs. "Since that time, Ms. Lavigne has refused..." (A) \u2014 functionally identical but B\u2019s is crisper.')
add_body(doc, "Document C restructures the Introduction into short stacked paragraphs, each advancing a discrete point: procedural posture, factual summary, legal conclusion, relief requested. This is more words overall but achieves greater compression per paragraph \u2014 each sentence does one job. C also adds a formal Conclusion section that A and B lack. That is standard in formal filings and not padding; its absence in A and B is actually the gap.")
add_edge(doc, "Edge: B for word economy in continuous prose. C for structural efficiency across the whole document.")

add_sub_heading(doc, "2. Plain Language vs. Unnecessary Legalese")
add_body(doc, 'All three documents are relatively restrained \u2014 no "heretofore," "inasmuch as," or similar affectations. The differences are in how much legal vocabulary does actual work.')
add_body(doc, 'Document A states the conversion definition as received formula \u2014 "the wrongful exercise of dominion over the property of another, to the exclusion of the owner\u2019s rights, or the withholding of possession under a claim of title inconsistent with the owner\u2019s" \u2014 and moves on. Technically precise but offers no elaboration on what that framework means in practice.')
add_body(doc, 'Document B uses the same definition verbatim. Its one small improvement: "the non-moving party may not rest on its pleadings" (B) vs. "cannot rest" (A) \u2014 "may not" more accurately reflects the procedural rule as a standard of law rather than mere incapacity.')
add_body(doc, 'Document C extends the conversion standard with the court\u2019s own language: "Vermont courts describe conversion as occurring where a defendant appropriates property to its own use and beneficial enjoyment, exercises dominion over it in defiance of the owner\u2019s rights, or withholds it from the owner under a claim inconsistent with the owner\u2019s title." This is still legal language, but it is tactical \u2014 you are giving the judge the framework she will likely apply. C also adds: "Courts applying Vermont law have granted partial summary judgment on conversion liability where the documentary record shows undisputed wrongful dominion." That sentence is plain, direct, and tells the court exactly what to do with the facts before it.')
add_edge(doc, "Edge: C for building in the court\u2019s own framework without resort to ornate legalese.")

add_sub_heading(doc, "3. Run-On Sentences")
add_body(doc, "This is the area of greatest variation across the three documents.")
add_body(doc, 'Document A\u2019s Introduction opens with a dense sentence: "As set forth below, the undisputed facts established by documentary evidence demonstrate that Respondents Jeanne Lavigne and Kate Remillard, acting without any authority as executor, trustee, or agent, seized the Firearms from the residence of the late Robert Marks and surrendered them to law enforcement." Three embedded phrases, two independent actions. Manageable, but the reader must work to extract the subject-verb-object chain. Later in Section B: "As the sole beneficiary with the vested right to receive this specific property, Petitioner holds the equitable title and the right to possession, subject only to proper administration of the estate and Trust \u2014 administration that Respondents have bypassed and obstructed." Not technically a run-on, but it requires three passes through the dependent structure.')
add_body(doc, 'Document B handles the equivalent of A\u2019s dense opener more elegantly using em-dashes: "Respondents Jeanne Lavigne and Kate Remillard \u2014 acting without authority as executor, trustee, or agent \u2014 entered the residence of the late Robert Marks, seized the Firearms, and surrendered them to law enforcement." The parenthetical is visually set off, and the main action ("entered... seized... surrendered") reads as a clear three-beat sequence. Same length as A, but easier to scan.')
add_body(doc, 'Document C avoids the problem structurally. What A and B compress into one or two dense sentences, C distributes across three short paragraphs: "In November 2024, Respondents Jeanne Lavigne and Kate Remillard entered the residence of Petitioner\u2019s late father, Robert Marks, and removed the Firearms." Clean. "They then surrendered the Firearms to law enforcement." Eight words. The deliberate short-sentence rhythm is the strongest in the set.')
add_edge(doc, "Edge: C. The Introduction structure avoids clause-chaining entirely. B\u2019s em-dash technique is an effective second-best.")

add_sub_heading(doc, "4. Repetition Across Sections")
add_body(doc, "Summary judgment motions inherently repeat \u2014 SUMF states the facts, the Argument reapplies them with analysis. That structural repetition is appropriate. What matters is whether any document repeats unnecessarily beyond that.")
add_body(doc, "Documents A and B are nearly identical in this regard. Both restate the three conversion elements in the analysis and again in the proposed declaration, which is expected and correct. Neither document introduces extra redundancy beyond the genre.")
add_body(doc, "Document C shows more structural control. Its SUMF splits Robert\u2019s death date (SUMF \u00b6 1) from his firearms ownership (SUMF \u00b6 2) into separate numbered items, creating 12 items vs. 11 in A and B. This is better practice \u2014 each fact stands alone, limiting the surface area for opposition argument about compound factual assertions.")
add_body(doc, "More importantly, C front-loads the full narrative into the Introduction so cleanly that the Argument section can proceed directly to element-by-element analysis without re-narrating. In A and B, the Introduction is compressed and the background story has to be reconstructed again in Section B. C avoids that doubling.")
add_edge(doc, "Edge: C for structural discipline.")

add_sub_heading(doc, "5. Active vs. Passive Voice")
add_body(doc, "Active voice matters most in a conversion claim, where naming who did what to whom is the core of the argument.")
add_body(doc, 'Document A is generally active but buries the wrongful actors in a participial phrase in the Introduction: "Respondents... acting without any authority... seized the Firearms" \u2014 the actors are present but not at the front of the sentence. The passive-adjacent phrase "her exercise of dominion was wrongful as a matter of law" is a standard legal formulation and not a meaningful weakness, but the Introduction sets a softer tone than it needs to.')
add_body(doc, 'Document B is structurally identical to A on voice but uses the em-dash construction to pull the wrongdoers forward more visually: "Respondents Jeanne Lavigne and Kate Remillard \u2014 acting without authority... \u2014 entered the residence... seized the Firearms, and surrendered them." The actors remain the subject and the actions are the verbs; B\u2019s punctuation makes it land harder than A\u2019s.')
add_body(doc, 'Document C is the clearest here. The Introduction gives us: "Respondents Jeanne Lavigne and Kate Remillard entered the residence of Petitioner\u2019s late father, Robert Marks, and removed the Firearms. They then surrendered the Firearms to law enforcement." Two short, active sentences. The wrongdoers are the grammatical subjects of every clause. In Section B-2, the pattern repeats: "Ms. Lavigne and Ms. Remillard entered the residence, took possession of the Firearms, and surrendered them to the Grand Isle County Sheriff\u2019s Office." The verbs are vivid and the actors own every one of them.')
add_edge(doc, "Edge: C. The Introduction especially \u2014 the most disciplined active-voice narration in the set, and most appropriate for a conversion claim.")

add_sub_heading(doc, "6. Clarity of Argument")
add_body(doc, "The test: can a smart non-lawyer read this and immediately understand what happened, why it is wrong, and what the writer wants the court to do?")
add_body(doc, 'Document A achieves this, but requires effort. The Introduction announces "undisputed facts" and "wrongful dominion" but does not give a reader the full picture \u2014 the absence of any probate proceeding, the death-terminated power of attorney, the trust chain \u2014 until after the SUMF and well into the Argument. A non-lawyer has to keep reading and cross-referencing to assemble the story.')
add_body(doc, "Document B is essentially the same as A in this regard \u2014 the prose is tighter, but the information architecture is identical.")
add_body(doc, 'Document C inverts the structure. By the end of the four-paragraph Introduction, a non-lawyer knows: (1) who the decedent was and his relationship to Petitioner; (2) exactly what the respondents did in November 2024; (3) that they had no legal authority whatsoever \u2014 three independent grounds; (4) that Petitioner is the sole beneficiary under the trust; and (5) that this is conversion as a matter of law. The legal conclusion appears in the Introduction itself: "Respondents\u2019 unauthorized seizure and continued withholding of property to which Petitioner has the right of possession constitutes conversion as a matter of law." The Argument section then proves what the Introduction already told the court to find.')
add_body(doc, 'C also uses numbered sub-headings within Section B (rather than the "First... Second... Third" prose structure of A and B), making the three-element test immediately scannable. A reader can locate any element directly.')
add_edge(doc, "Edge: C, by a meaningful margin.")

add_sub_heading(doc, "7. Overall Persuasive Force")
add_body(doc, 'Document A is a competent, well-organized motion. The facts are correctly assembled, the legal standard is properly applied, and the conclusion follows. But it feels like it is executing a checklist rather than pressing a case. The phrase "As set forth below" in the first sentence of the Introduction is symptomatic \u2014 it is procedural throat-clearing that signals the document is thinking about its own structure rather than the judge it is trying to persuade.')
add_body(doc, "Document B is Document A after disciplined editorial attention. It is tighter in nearly every instance the two diverge. The em-dash technique in the Introduction is the most visible stylistic signature and it works. But B shares A\u2019s fundamental posture: it reports the wrong methodically and correctly rather than presenting it as an argument that wants to be won.")
add_body(doc, 'Document C has more life. The phrase "Petitioner\u2019s late father, Robert Marks" is a small but intentional choice \u2014 it reminds the court that this is a family situation involving grief and a son trying to claim his father\u2019s property. It does not manipulate; it contextualizes. The short staccato sentences of the Introduction give the wrongdoing rhetorical force \u2014 they make the sequence of acts feel deliberate and concrete. The expanded proposed declaration at the end of Section C tells the court exactly what to order across three numbered sub-clauses, reducing judicial workload and leaving less room for ambiguity in the order. The Vermont case citations signal preparation.')
add_edge(doc, "Edge: C for overall persuasive force.")

# ---- SECTION 2 ----
add_section_heading(doc, "Holistic Summary")

add_sub_heading(doc, "Document A")
add_body(doc, "A solid, professional summary judgment motion that would serve its purpose without embarrassment. The legal analysis is correct, the coverage is complete, and the prose is clear. But it reads as a capable draft \u2014 organized, accurate, and methodical \u2014 rather than as a piece of advocacy. The filler opening phrase, the double-header on the WHEREFORE section, and the compressed Introduction give it the quality of a document that organizes facts well without leaning into the argument. A judge would read it and understand the claim. A judge would not feel the urgency of the wrong.")

add_sub_heading(doc, "Document B")
add_body(doc, 'Document B is Document A after disciplined editorial attention. It is tighter in nearly every place the two differ. The em-dash technique in the Introduction is the most visible stylistic signature and it works well. Sentence-level choices are consistently better \u2014 "may not rest," "direct exercise of dominion," cleaner SUMF punctuation. But B shares A\u2019s fundamental restraint: it reports rather than presses. If B and A were the only choices, B is the better filing without hesitation \u2014 and by more than the small number of differences suggests.')

add_sub_heading(doc, "Document C")
add_body(doc, 'Document C treats the same content as a persuasive instrument rather than an informational record. The restructured Introduction, numbered sub-headings, formal Conclusion, and case citations all reflect a document that has been architected with the reader\u2019s experience in mind. The humanizing phrase about Robert Marks as "Petitioner\u2019s late father," the short active sentences, the detailed proposed declaration \u2014 these are the marks of a drafter who understood that a judge reading a motion for partial summary judgment wants to reach the right answer quickly, and this document makes that path as short as possible. C is the most fully realized of the three as a piece of advocacy.')

add_sub_heading(doc, "Cross-Document Synthesis")
add_body(doc, "Documents A and B are closely related and share their informational architecture entirely; B is the more refined version of A. C is a different approach \u2014 it rebuilds the Introduction, imposes structural headings, adds citations, and chooses its details deliberately. The pattern suggests A was a first-generation or template-based output, B applied careful editorial judgment to A, and C approached the task with more deliberate structural thinking from the start. All three would achieve the same legal result. C would achieve it while leaving a more favorable impression on any reader who has to decide.")

# ---- SECTION 3 ----
add_section_heading(doc, "Situational Preference")

add_body(doc, "For a Vermont trial court in a probate/civil division: Document C is preferred. The Vermont-specific citations (Price v. Leland, Berlin Dev. Assocs.) signal preparation and give the court ready Vermont authority to cite in any written ruling. Trial courts appreciate motions that do this work for them.")
add_body(doc, "For a strict formal record or appellate court: Document B is the most polished in terms of consistent formatting, punctuation discipline, and absence of ambiguous structural choices (A\u2019s double WHEREFORE label would need to go; C\u2019s exhibit reference discrepancy \u2014 Will cited as Exhibit G in C vs. Exhibit C in A/B \u2014 would need to be resolved before filing). B and C are both appropriate; B requires less pre-filing cleanup.")
add_body(doc, "For a judge known to prefer dense, economical briefs: Document B is the most efficient. Document A is competitive. Document C\u2019s additional structure and Conclusion section add length.")
add_body(doc, "For maximum persuasive impact in any context: Document C, provided the exhibit citation discrepancy is confirmed and the case citations verified. The active-voice Introduction, the humanizing detail, the structured argument, and the precise proposed declaration all work together.")
add_body(doc, "If starting from one of these to file today: Document B is the safest ready-to-file choice \u2014 cleaner than A, shares the same verified structure, requires no exhibit reconciliation. Document C is the best choice if the goal is a fully crafted advocacy document with the additional polish applied to resolve the exhibit issue.")

add_summary_table(doc, [
    ("Conciseness", "B overall; C in sections"),
    ("Plain language", "C builds on court\u2019s language"),
    ("Run-on sentences", "C clearly"),
    ("Repetition", "C for structural discipline"),
    ("Active voice", "C clearly"),
    ("Clarity of argument", "C clearly; A weakest"),
    ("Persuasive force", "C overall"),
])

doc.save(output_path)
print(f"Saved: {output_path}")
